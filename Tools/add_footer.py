#!/usr/bin/env python3
"""
add_footer.py
Adds a footer to a .docx file:
  - Left:  filename at 9pt
  - Right: current page / total pages at 9pt

Usage:
    python3 add_footer.py <docx_file>

The script manages its own virtual environment at /tmp/docenv.
"""

import sys
import os
import subprocess

VENV_DIR = os.path.join(os.path.expanduser("~"), ".docenv")
VENV_PYTHON = os.path.join(VENV_DIR, "bin", "python3")
VENV_PIP = os.path.join(VENV_DIR, "bin", "pip")


def is_venv_ready():
    if not os.path.exists(VENV_PYTHON):
        return False
    result = subprocess.run(
        [VENV_PYTHON, "-c", "import docx"],
        capture_output=True
    )
    return result.returncode == 0


def setup_venv():
    print("Setting up document-processing environment...", file=sys.stderr)
    subprocess.run([sys.executable, "-m", "venv", VENV_DIR], check=True)
    subprocess.run([VENV_PIP, "install", "python-docx"],
                   check=True, stdout=subprocess.DEVNULL)
    print("Setup complete.", file=sys.stderr)


if not os.path.realpath(sys.prefix).startswith(os.path.realpath(VENV_DIR)):
    if not is_venv_ready():
        setup_venv()
    os.execv(VENV_PYTHON, [VENV_PYTHON] + sys.argv)


# ── Running inside venv ───────────────────────────────────────────────────────

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_field(run, field_name):
    """Insert a Word field (e.g. PAGE, NUMPAGES) into a run."""
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_name} '
    run._r.append(instr)

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.append(end)


def set_right_tab(para, pos_twips):
    """Add a right-aligned tab stop to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    # Remove any existing tabs element
    for existing in pPr.findall(qn('w:tabs')):
        pPr.remove(existing)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def add_footer(docx_path):
    filename = os.path.basename(docx_path)
    doc = Document(docx_path)
    section = doc.sections[0]

    # Calculate text width in twips for right tab stop
    # python-docx returns measurements in EMU (English Metric Units)
    # 1 inch = 914400 EMU = 1440 twips, so 1 twip = 635 EMU
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    text_width_twips = int((page_width - left_margin - right_margin) / 635)

    footer = section.footer
    section.different_first_page_header_footer = False

    # Use existing paragraph or create one
    if footer.paragraphs:
        para = footer.paragraphs[0]
        para.clear()
    else:
        para = footer.add_paragraph()

    # Set right tab at text width
    set_right_tab(para, text_width_twips)

    # Left: filename at 9pt
    run_left = para.add_run(filename)
    run_left.font.size = Pt(9)

    # Tab to right
    run_tab = para.add_run('\t')
    run_tab.font.size = Pt(9)

    # Right: PAGE / NUMPAGES at 9pt
    run_page = para.add_run()
    run_page.font.size = Pt(9)
    add_field(run_page, 'PAGE')

    run_sep = para.add_run(' / ')
    run_sep.font.size = Pt(9)

    run_total = para.add_run()
    run_total.font.size = Pt(9)
    add_field(run_total, 'NUMPAGES')

    doc.save(docx_path)
    print(f"Footer added to: {docx_path}", file=sys.stderr)


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 add_footer.py <docx_file>", file=sys.stderr)
        sys.exit(1)
    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"Error: file not found: {docx_path}", file=sys.stderr)
        sys.exit(1)
    add_footer(docx_path)


if __name__ == "__main__":
    main()
